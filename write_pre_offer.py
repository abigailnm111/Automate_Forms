
"""
Created on Tue Apr 13 10:24:12 2021

@author: panda
"""

from datetime import date
import re
from docx import Document, shared
import docx2txt

#pre_template= docx2txt.process('Pre6-letter-template.docx')
    #template for continuing lecturers
#cont_template=docx2txt.process('12_CL_yearly_assign_letter.docx')

#print(cont_template)

today= date.today()
letter_date= today.strftime("%B %d, %Y")

#formats courses to clearly state how many sections of each course are assigned. 
#ignores course releases
def format_course(course):
    for c in course:
            title= course.split(' ',1)
            if ("course" in title[1]) or ("Course" in title[1]):
                course_format= course
            else:   
                course_format= title[0]+" section(s): "+title[1]
            
    return course_format   

#formats courses into groups by quarters with formatting between each quarter of courses
def format_courses_by_quarter(lecturer):
    all_courses=[]
    if lecturer.F_courses!=[]:
        c= "".join(str(format_course(course)+",\n\t\t\t\t") for course in lecturer.F_courses)
        courses=c[:-6]+c[-5:]
        all_courses.append("Fall: "+courses)
        
    if lecturer.W_courses!=[]:
      
        c= "".join(str(format_course(course)+",\n\t\t\t\t") for course in lecturer.W_courses)
        courses=c[:-6]+c[-5:]
        all_courses.append("Winter: "+courses)
        
    if lecturer.S_courses!=[]:
        
        c= "".join(str(format_course(course)+",\n\t\t\t\t") for course in lecturer.S_courses)
        courses=c[:-6]+c[-5:]
        all_courses.append("Spring: "+courses)

    courses= "".join(str(quarter) for quarter in all_courses)
    return courses

def write_letter(lecturer, AY):
    #template for pre 6 lecturers
    pre_template= docx2txt.process('Pre6-letter-template.docx')
    #template for continuing lecturers
    cont_template=docx2txt.process('12_CL_yearly_assign_letter.docx')
    
    #letterhead for letters
    new_doc=Document('Letterhead.docx')
    #formats course listings for letter by quarter
    courses=format_courses_by_quarter(lecturer)
    #chooses template by title code
    if lecturer.job_code[0]==1630 or lecturer.job_code[0]== 1632:
        template= pre_template
    else:
        template= cont_template
    AY=AY+"-"+str(int(AY)+1)
    job_code=str(lecturer.job_code[0])

    i=0
    dates=' '
    #iterates through start dates to format data for letter
    for s_date in lecturer.start:
            start_end= lecturer.start[i]+" - "+lecturer.end[i]
            annual=lecturer.annual[0]
            monthly=lecturer.monthly[0]
            #accounts for multiple start dates
            if i== 1:
                start_end= ", "+start_end
                #if lecturer goes from one job code to another (initial continuing or sr.), it will not add second start date
                if len(job_code)==1:
                    if len(lecturer.annual)>1:
                        if lecturer.annual[1]== "ATTENTION":
                            annual= annual+", "+lecturer.start[i]+': '+lecturer.annual[1]  
                            monthly=monthly+", "+lecturer.start[i]+': '+lecturer.monthly[1]   
                else:
                    start_end=""
            dates= dates+start_end
            i+=1
            job_code=str(lecturer.job_code[0])
    #dictionary to iterage through RegEx on template
    data={"<Date>":str(letter_date) ,"<Name>":lecturer.first_name+" "+lecturer.last_name, 
          "<Title Code>": job_code, "<start to end>": dates, 
          "<percentage>":lecturer.percentage, "<annual>":"{:,}".format(annual), "<monthly>": "{:,.2f}".format(monthly),
          "<List of Classes & Course Titles>": courses, "\n\n": '\n', "<AY>":AY}
    #replace keywords with formatted data
    for key in data:
        template=re.sub(key, data[key] , template)
    #Save letters
    run=new_doc.add_paragraph(template). add_run()
    style=new_doc.styles['Normal']
    font=style.font
    font.name=  "Garamond"
    font.size= shared.Pt(11)
    new_doc.save(AY+lecturer.last_name+"_"+lecturer.first_name+'.docx')